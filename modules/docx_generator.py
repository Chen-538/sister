"""DOCX 產生模組 - 讀取範本並填入資料"""

import re
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

TEMPLATES_DIR = Path(__file__).resolve().parent.parent / "templates"


def _set_cell_text(cell, text):
    """清除儲存格原有內容，寫入新文字，保留第一段的格式。"""
    if isinstance(text, list):
        text = "\n".join(str(item) for item in text)
    text = str(text)

    # 取得第一段的段落格式與字型格式
    first_para = cell.paragraphs[0]
    para_fmt = first_para.paragraph_format
    font_props = {}
    if first_para.runs:
        src_font = first_para.runs[0].font
        if src_font.size:
            font_props["size"] = src_font.size
        if src_font.name:
            font_props["name"] = src_font.name
        if src_font.bold is not None:
            font_props["bold"] = src_font.bold

    # 刪除儲存格內所有段落 XML（除了第一段）
    tc = cell._tc
    for p_elem in tc.findall(qn("w:p"))[1:]:
        tc.remove(p_elem)

    # 清除第一段所有 run
    for run in first_para.runs:
        first_para._p.remove(run._r)

    # 寫入新內容，每個 \n 換一段
    lines = text.split("\n")
    for i, line in enumerate(lines):
        if i == 0:
            para = first_para
        else:
            # 新增段落：複製段落格式 XML
            new_p = tc.makeelement(qn("w:p"), {})
            # 複製段落屬性
            src_pPr = first_para._p.find(qn("w:pPr"))
            if src_pPr is not None:
                from copy import deepcopy
                new_p.append(deepcopy(src_pPr))
            tc.append(new_p)
            from docx.text.paragraph import Paragraph
            para = Paragraph(new_p, cell)

        run = para.add_run(line)
        for prop, val in font_props.items():
            setattr(run.font, prop, val)


def _set_para_text(para, new_text: str):
    """替換整個段落的文字，保留第一個 run 的格式。"""
    font_props = {}
    if para.runs:
        src_font = para.runs[0].font
        if src_font.size:
            font_props["size"] = src_font.size
        if src_font.name:
            font_props["name"] = src_font.name
        if src_font.bold is not None:
            font_props["bold"] = src_font.bold

    # 刪除所有 run
    for run in para.runs:
        para._p.remove(run._r)

    # 寫入新 run
    run = para.add_run(new_text)
    for prop, val in font_props.items():
        setattr(run.font, prop, val)


def generate_newsletter(structured_data: dict, newsletter_content: dict, output_path: str) -> str:
    """
    生成班刊 DOCX。

    班刊範本表格結構（5行 x 3列）：
    R0: 主題名稱 | (主題) | 週次日期
    R1: 活動名稱 | (活動) | (日期+週次)
    R2: 教學活動 | (內容, span=2)
    R3: 老師的叮嚀 | (內容, span=2)
    R4: 給老師的話 | (空白, span=2)
    """
    doc = Document(str(TEMPLATES_DIR / "班刊範本.docx"))

    # 更新標題段落中的出刊日期
    for para in doc.paragraphs:
        if "出刊" in para.text:
            出刊日期 = structured_data.get("出刊日期", "[待補]")
            _set_para_text(para, f"{出刊日期}出刊")

    table = doc.tables[0]

    # R0C1: 主題名稱
    _set_cell_text(table.rows[0].cells[1], structured_data.get("主題名稱", "[待補]"))

    # R1C1: 活動名稱
    _set_cell_text(table.rows[1].cells[1], structured_data.get("活動名稱", "[待補]"))
    # R1C2: 日期+週次
    日期區間 = structured_data.get("日期區間", "[待補]")
    週次 = structured_data.get("週次", "[待補]")
    _set_cell_text(table.rows[1].cells[2], f"{日期區間}\n第{週次}週")

    # R2C1: 教學活動（合併欄位，寫入 C1 即可）
    教學活動 = newsletter_content.get("教學活動", "[待補]")
    _set_cell_text(table.rows[2].cells[1], 教學活動)

    # R3C1: 老師叮嚀
    老師叮嚀 = newsletter_content.get("老師叮嚀", "[待補]")
    _set_cell_text(table.rows[3].cells[1], 老師叮嚀)

    # R4: 給老師的話 — 留空（家長填寫）

    doc.save(output_path)
    return output_path


def generate_weekly_log(structured_data: dict, log_content: dict, output_path: str) -> str:
    """
    生成週誌 DOCX。

    週誌範本表格結構（11行 x 4列）：
    R0:  日期（span=4）
    R1:  活動概要(span=3) | 學習指標
    R2:  活動概要內容(span=3) | 學習指標內容
    R3:  教學省思（span=4，標題列）
    R4:  教學省思內容（span=4）
    R5:  軼事記錄（span=4，標題列）
    R6:  類別 | 日期 | 情況描述(span=2)
    R7:  幼兒行為輔導 | 日期 | 描述(span=2)
    R8:  親師溝通 | 日期 | 描述(span=2)
    R9:  照片記錄（span=4，標題列）
    R10: 照片記錄內容（span=4）
    """
    doc = Document(str(TEMPLATES_DIR / "週誌範本.docx"))

    # 更新標題段落
    for para in doc.paragraphs:
        text = para.text.strip()
        if "週教學週誌" in text:
            週次 = structured_data.get("週次", "[待補]")
            new_text = re.sub(r"第\d+週", f"第{週次}週", para.text)
            _set_para_text(para, new_text)
        elif "主題名稱" in text:
            主題 = structured_data.get("主題名稱", "[待補]")
            教師 = structured_data.get("教師姓名", "[待補]")
            _set_para_text(para, f"主題名稱：{主題}                     教師：{教師}")

    table = doc.tables[0]

    # R0: 日期
    日期區間 = structured_data.get("日期區間", "[待補]")
    _set_cell_text(table.rows[0].cells[0], f"日期： {日期區間}")

    # R2C0: 活動概要內容
    活動概要 = log_content.get("活動概要", "[待補]")
    _set_cell_text(table.rows[2].cells[0], 活動概要)

    # R2C3: 學習指標
    學習指標 = log_content.get("學習指標", [])
    if isinstance(學習指標, list):
        學習指標_text = "\n\n".join(學習指標)
    else:
        學習指標_text = str(學習指標)
    _set_cell_text(table.rows[2].cells[3], 學習指標_text)

    # R4: 教學省思內容
    教學省思 = log_content.get("教學省思", "[待補]")
    _set_cell_text(table.rows[4].cells[0], f"教學上的想法及觀察\n{教學省思}")

    # R7: 行為輔導
    行為輔導 = log_content.get("行為輔導", {})
    if isinstance(行為輔導, dict):
        _set_cell_text(table.rows[7].cells[1], 行為輔導.get("日期", "[待補]"))
        _set_cell_text(table.rows[7].cells[2], 行為輔導.get("描述", "[無]"))
    else:
        _set_cell_text(table.rows[7].cells[2], str(行為輔導))

    # R8: 親師溝通
    親師溝通 = log_content.get("親師溝通", {})
    if isinstance(親師溝通, dict):
        _set_cell_text(table.rows[8].cells[1], 親師溝通.get("日期", "[待補]"))
        _set_cell_text(table.rows[8].cells[2], 親師溝通.get("描述", "[無]"))
    else:
        _set_cell_text(table.rows[8].cells[2], str(親師溝通))

    # R10: 照片記錄說明
    照片說明 = log_content.get("照片記錄說明", "[無]")
    _set_cell_text(table.rows[10].cells[0], 照片說明)

    doc.save(output_path)
    return output_path
